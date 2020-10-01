import docx, os

file = docx.Document(os.path.dirname(__file__)+"/"+"Oleinik.docx")

print(len(file.paragraphs))

file.add_paragraph("Hello world!")



for i in range(len(file.paragraphs)):
    print(file.paragraphs[i].text)
    file.paragraphs[i].add_run("qqqq")

file.add_picture("filename.bmp", width=docx.shared.Inches(1) , height=docx.shared.Cm(5))
    

print(len(file.paragraphs[14].runs))

print(file.paragraphs[14].text)

for i in range(len(file.paragraphs[14].runs)):
    print(f"Runs #{i} -", file.paragraphs[14].runs[i].text)

file.save(os.path.dirname(__file__)+"/"+"oleinik1.docx")




